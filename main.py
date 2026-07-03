from fastapi import FastAPI
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from agent.orchestrator_simple import run_orchestrator

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

global charter_text_global, wbs_global, plan_global
charter_text_global = None
wbs_global = None
plan_global = None

# Serve static files
try:
    app.mount("/static", StaticFiles(directory="ui"), name="static")
except:
    pass

# Serve chat.html
@app.get("/")
async def root():
    return FileResponse("ui/chat.html")

# Chat endpoint
@app.post("/chat")
async def chat(request: dict):
    """
    Chat endpoint - routes to orchestrator.
    
    Request: {
        "message": user message,
        "charter": charter text (optional),
        "wbs": WBS text (optional),
        "budget": budget (optional)
    }
    """
    try:
        message = request.get("message", "").strip()
        #charter = request.get("charter")
        #wbs = request.get("wbs")
        #budget = request.get("budget", 250000)
        
        if not message:
            return JSONResponse({"success": False, "error": "Empty message"}, status_code=400)
        
        # Call orchestrator
        result = run_orchestrator(message,budget=250000)
        
        if result.get("success"):
            return JSONResponse({
                "success": True,
                "type": result.get("type"),
                "content": result.get("content")
            })
        else:
            return JSONResponse({
                "success": False,
                "error": result.get("error")
            }, status_code=400)
    
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({
            "error": str(e),
            "success": False
        }, status_code=500)

# Export endpoint
@app.post("/export")
async def export(request: dict):
    """
    Export to Word document.
    
    Request: {
        "type": "charter" | "wbs" | "plan",
        "content": content to export
    }
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return JSONResponse({
            "success": False,
            "error": "python-docx not installed. Run: pip install python-docx"
        }, status_code=400)
    
    try:
        doc_type = request.get("type", "document")
        content = request.get("content", "")
        
        if not content:
            return JSONResponse({
                "success": False,
                "error": "No content to export"
            }, status_code=400)
        
        # Create document
        doc = Document()
        
        # Add title
        titles = {
            'charter': 'Project Charter',
            'wbs': 'Work Breakdown Structure',
            'plan': 'Detailed Project Plan'
        }
        
        title = doc.add_heading(titles.get(doc_type, 'Document'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        metadata.add_run(f'Document Type: {titles.get(doc_type, "Document")}')
        metadata_format = metadata.paragraph_format
        metadata_format.space_before = Pt(12)
        metadata_format.space_after = Pt(12)
        
        # Add content
        content_para = doc.add_paragraph(content)
        content_para.paragraph_format.space_before = Pt(12)
        
        # Save to temp file
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Return file
        response = FileResponse(
            tmp_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f'{doc_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
        
        # Clean up after response
        import atexit
        atexit.register(lambda: os.unlink(tmp_path) if os.path.exists(tmp_path) else None)
        
        return response
    
    except Exception as e:
        print(f"Export error: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({
            "success": False,
            "error": str(e)
        }, status_code=500)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)